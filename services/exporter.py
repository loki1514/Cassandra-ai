"""
services/exporter.py — Meeting DOCX Export

Fetches all transcripts and artifacts for a given meeting from Supabase
and produces a structured .docx file using python-docx.
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from infrastructure.database import get_supabase

supabase = get_supabase()


def _heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def generate_meeting_docx(meeting_id: str) -> bytes:
    """
    Builds a professional DOCX for the given meeting_id.
    Returns the file as raw bytes.
    """
    doc = Document()

    # ── Style tweaks ──────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Title block ───────────────────────────────────────────
    title = doc.add_heading("MEETING MINUTES", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Meeting ID: {meeting_id}").bold = True
    doc.add_paragraph(
        f"Generated: {datetime.utcnow().strftime('%B %d, %Y  %H:%M UTC')}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    # ── Fetch transcripts ─────────────────────────────────────
    tx_result = (
        supabase.table("transcripts")
        .select("speaker_role, content, chunk_index, created_at")
        .eq("meeting_id", meeting_id)
        .order("chunk_index")
        .execute()
    )
    transcripts = tx_result.data or []

    # ── Fetch artifacts ───────────────────────────────────────
    art_result = (
        supabase.table("artifacts")
        .select("artifact_type, content, confidence, created_at")
        .eq("meeting_id", meeting_id)
        .order("created_at")
        .execute()
    )
    artifacts = art_result.data or []

    decisions = [a for a in artifacts if a["artifact_type"] == "decision"]
    risks     = [a for a in artifacts if a["artifact_type"] == "risk"]
    topics    = [a for a in artifacts if a["artifact_type"] == "topic"]
    summaries = [a for a in artifacts if a["artifact_type"] == "summary"]

    # ── TRANSCRIPT SECTION ────────────────────────────────────
    _heading(doc, "FULL TRANSCRIPT", level=1)
    if transcripts:
        for row in transcripts:
            speaker = row.get("speaker_role", "unknown").upper()
            content = row.get("content", "").strip()
            if not content:
                continue
            p = doc.add_paragraph()
            run_label = p.add_run(f"{speaker}: ")
            run_label.bold = True
            p.add_run(content)
    else:
        doc.add_paragraph("No transcript available for this meeting.")

    doc.add_page_break()

    # ── KEY DECISIONS ─────────────────────────────────────────
    _heading(doc, "KEY DECISIONS", level=1)
    if decisions:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Decision"
        hdr[1].text = "Confidence"
        for hc in hdr:
            for para in hc.paragraphs:
                for run in para.runs:
                    run.bold = True
        for d in decisions:
            row_cells = table.add_row().cells
            row_cells[0].text = d.get("content", "")
            conf = d.get("confidence", 0.9)
            row_cells[1].text = f"{int(conf * 100)}%" if conf <= 1 else f"{conf}%"
    else:
        doc.add_paragraph("No decisions recorded.")

    doc.add_paragraph()

    # ── RISK FLAGS ────────────────────────────────────────────
    _heading(doc, "RISK FLAGS", level=1)
    if risks:
        for r in risks:
            doc.add_paragraph(r.get("content", ""), style="List Bullet")
    else:
        doc.add_paragraph("No risks recorded.")

    doc.add_paragraph()

    # ── ACTION ITEMS (topics) ─────────────────────────────────
    _heading(doc, "ACTION ITEMS", level=1)
    if topics:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Action"
        hdr[1].text = "Notes"
        for hc in hdr:
            for para in hc.paragraphs:
                for run in para.runs:
                    run.bold = True
        for t in topics:
            row_cells = table.add_row().cells
            row_cells[0].text = t.get("content", "")
            row_cells[1].text = ""  # Placeholder for owner/due date
    else:
        doc.add_paragraph("No action items recorded.")

    doc.add_paragraph()

    # ── SUMMARIES ─────────────────────────────────────────────
    if summaries:
        _heading(doc, "SUMMARIES", level=1)
        for s in summaries:
            doc.add_paragraph(s.get("content", ""), style="List Bullet")

    # ── Serialize ─────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
